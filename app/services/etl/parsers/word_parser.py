"""Word document parser for `.docx` files."""

from __future__ import annotations

from pathlib import Path
from typing import Mapping
import logging
import uuid
import xml.etree.ElementTree as ET
import zipfile

from app.services.etl.models import ExtractedSegment, ParsedDocument
from app.services.etl.utils import normalize_text_block

try:
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover - fallback path only used when python-docx is unavailable.
    DocxDocument = None


_WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
logger = logging.getLogger(__name__)


class WordDocumentParser:
    """Parse `.docx` files into normalized Aurora ETL documents."""

    parser_name = "word_document_parser"
    supported_extensions = {".docx"}

    def parse(
        self,
        *,
        file_path: Path,
        relative_path: str,
        extra_metadata: Mapping[str, object] | None = None,
    ) -> ParsedDocument:
        """Parse one `.docx` file."""
        blocks = self._extract_blocks(file_path)
        normalized_blocks = [normalize_text_block(block) for block in blocks if normalize_text_block(block)]
        if not normalized_blocks:
            raise ValueError(f"Document is empty after parsing: {file_path.name}")

        source_id = uuid.uuid4().hex
        markdown_body = "\n\n".join(normalized_blocks)
        markdown_text = f"# {file_path.name}\n\n{markdown_body}".strip()
        segment = ExtractedSegment(
            segment_id=uuid.uuid4().hex,
            sequence=1,
            content_text=markdown_body,
            content_markdown=markdown_text,
            metadata={
                "segment_kind": "document",
                "source_type": "docx",
            },
        )
        content_json = {
            "source_id": source_id,
            "file_name": file_path.name,
            "file_type": "docx",
            "parser_name": self.parser_name,
            "segment_count": 1,
            "segments": [segment.to_dict()],
        }
        return ParsedDocument(
            source_id=source_id,
            source_path=str(file_path),
            relative_path=relative_path,
            file_name=file_path.name,
            file_type="docx",
            parser_name=self.parser_name,
            content_markdown=markdown_text,
            content_json=content_json,
            segments=[segment],
            metadata=dict(extra_metadata or {}),
        )

    def _extract_blocks(self, file_path: Path) -> list[str]:
        if DocxDocument is not None:
            try:
                return self._extract_with_python_docx(file_path)
            except Exception:
                logger.exception("python-docx parsing failed for %s, falling back to zip/xml.", file_path)
        return self._extract_with_zip(file_path)

    def _extract_with_python_docx(self, file_path: Path) -> list[str]:
        document = DocxDocument(str(file_path))
        blocks: list[str] = []

        for paragraph in document.paragraphs:
            text = normalize_text_block(paragraph.text)
            if text:
                blocks.append(text)

        for index, table in enumerate(document.tables, start=1):
            rows: list[list[str]] = []
            for row in table.rows:
                values = [normalize_text_block(cell.text) for cell in row.cells]
                if any(values):
                    rows.append(values)
            if rows:
                blocks.append(f"## Table {index}\n\n{_rows_to_markdown(rows)}")

        return blocks

    def _extract_with_zip(self, file_path: Path) -> list[str]:
        blocks: list[str] = []
        with zipfile.ZipFile(file_path) as archive:
            xml_bytes = archive.read("word/document.xml")

        root = ET.fromstring(xml_bytes)
        body = root.find("w:body", _WORD_NS)
        if body is None:
            return blocks

        for child in body:
            if child.tag.endswith("}p"):
                text = normalize_text_block("".join(node.text or "" for node in child.findall(".//w:t", _WORD_NS)))
                if text:
                    blocks.append(text)
                continue

            if child.tag.endswith("}tbl"):
                rows: list[list[str]] = []
                for row_node in child.findall(".//w:tr", _WORD_NS):
                    row_values: list[str] = []
                    for cell_node in row_node.findall(".//w:tc", _WORD_NS):
                        cell_text = normalize_text_block(
                            "".join(node.text or "" for node in cell_node.findall(".//w:t", _WORD_NS))
                        )
                        row_values.append(cell_text)
                    if any(row_values):
                        rows.append(row_values)
                if rows:
                    blocks.append(f"## Table {len([item for item in blocks if item.startswith('## Table ')]) + 1}\n\n{_rows_to_markdown(rows)}")

        return blocks


def _rows_to_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""

    max_columns = max(len(row) for row in rows)
    padded_rows = [row + [""] * (max_columns - len(row)) for row in rows]
    header = padded_rows[0]
    body = padded_rows[1:] or [[""] * max_columns]

    markdown_lines = [
        f"| {' | '.join(_escape_markdown_cell(cell) for cell in header)} |",
        f"| {' | '.join('---' for _ in range(max_columns))} |",
    ]
    markdown_lines.extend(
        f"| {' | '.join(_escape_markdown_cell(cell) for cell in row)} |"
        for row in body
    )
    return "\n".join(markdown_lines)


def _escape_markdown_cell(value: str) -> str:
    return str(value or "").replace("|", "\\|").replace("\n", " ").strip()
